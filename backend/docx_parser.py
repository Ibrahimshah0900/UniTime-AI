from pathlib import Path
import re

from docx import Document


DAY_MAP = {
    "MON": "Monday",
    "TUS": "Tuesday",
    "TUE": "Tuesday",
    "WED": "Wednesday",
    "THR": "Thursday",
    "THU": "Thursday",
    "FRI": "Friday",
}


TIME_RANGE_PATTERN = re.compile(
    r"^\s*\d{1,2}:\d{2}\s*(?:am|pm)?\s*-\s*"
    r"\d{1,2}:\d{2}\s*(?:am|pm)?\s*$",
    re.IGNORECASE,
)

TIME_TOKEN_PATTERN = re.compile(
    r"\b\d{1,2}:\d{2}\s*(?:am|pm)?\b",
    re.IGNORECASE,
)


def clean_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = value.replace("–", "-")
    value = value.replace("—", "-")

    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\s*\n\s*", "\n", value)

    return value.strip()


def is_header_like(value: str) -> bool:
    value = clean_text(value)

    if not value:
        return True

    if value.lower() == "break":
        return True

    if TIME_RANGE_PATTERN.match(value):
        return True

    # Handles malformed headers such as:
    # "08:30am 10:00am"
    return len(TIME_TOKEN_PATTERN.findall(value)) >= 2


def split_cell_entries(value: str) -> list[str]:
    value = clean_text(value)

    if not value:
        return []

    return [
        clean_text(line)
        for line in value.split("\n")
        if clean_text(line)
    ]


def extract_raw_timetable_records(
    file_path: str,
) -> tuple[list[dict], list[dict]]:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {path}"
        )

    document = Document(path)

    if not document.tables:
        raise ValueError(
            "No timetable tables found in the DOCX file."
        )

    table = document.tables[0]

    current_day: str | None = None
    current_slots: dict[int, str] = {}

    records: list[dict] = []
    warnings: list[dict] = []

    seen_records: set[
        tuple[str, str, str]
    ] = set()

    for row_number, row in enumerate(
        table.rows,
        start=1,
    ):
        raw_cells = [
            cell.text
            for cell in row.cells
        ]

        cells = [
            clean_text(value)
            for value in raw_cells
        ]

        if row_number == 1:
            continue

        raw_day = cells[0].upper()

        possible_headers = cells[1:]

        non_empty_headers = [
            value
            for value in possible_headers
            if value
        ]

        # A real day-header row should contain several
        # values that all look like times/breaks.
        is_day_header = (
            raw_day in DAY_MAP
            and len(non_empty_headers) >= 3
            and all(
                is_header_like(value)
                for value in non_empty_headers
            )
        )

        if is_day_header:
            current_day = DAY_MAP[raw_day]
            current_slots = {}

            for column_index, header in enumerate(
                possible_headers,
                start=1,
            ):
                header = clean_text(header)

                if not header:
                    continue

                if header.lower() == "break":
                    continue

                current_slots[column_index] = (
                    header.lower()
                )

                if not TIME_RANGE_PATTERN.match(
                    header
                ):
                    warnings.append(
                        {
                            "row": row_number,
                            "column": column_index + 1,
                            "type": "malformed_time_header",
                            "value": header,
                            "message": (
                                "Time header does not use the "
                                "expected start-end format."
                            ),
                        }
                    )

                if (
                    current_day == "Monday"
                    and "11:30pm"
                    in header.lower()
                ):
                    warnings.append(
                        {
                            "row": row_number,
                            "column": column_index + 1,
                            "type": "suspicious_time_header",
                            "value": header,
                            "message": (
                                "The source timetable says "
                                "'11:30pm - 01:00pm'. "
                                "The value has been preserved."
                            ),
                        }
                    )

            continue

        if current_day is None:
            continue

        for column_index, raw_cell in enumerate(
            raw_cells[1:],
            start=1,
        ):
            if column_index not in current_slots:
                continue

            time_slot = current_slots[
                column_index
            ]

            entries = split_cell_entries(
                raw_cell
            )

            for entry in entries:
                key = (
                    current_day,
                    time_slot,
                    entry,
                )

                # Word merged cells can expose the same
                # content several times.
                if key in seen_records:
                    continue

                seen_records.add(key)

                records.append(
                    {
                        "day": current_day,
                        "time_slot": time_slot,
                        "raw_text": entry,
                        "source_row": row_number,
                        "source_column": (
                            column_index + 1
                        ),
                    }
                )

    return records, warnings


def print_summary(
    records: list[dict],
    warnings: list[dict],
) -> None:
    print(
        f"Unique raw records: {len(records)}"
    )

    print(
        f"Warnings: {len(warnings)}"
    )

    counts: dict[str, int] = {}

    for record in records:
        day = record["day"]

        counts[day] = (
            counts.get(day, 0) + 1
        )

    print()
    print("Records by day:")

    for day in (
        "Monday",
        "Tuesday",
        "Wednesday",
        "Thursday",
        "Friday",
    ):
        print(
            f"{day}: {counts.get(day, 0)}"
        )

    print()
    print("First 25 records:")
    print("-" * 80)

    for record in records[:25]:
        print(
            f"{record['day']} | "
            f"{record['time_slot']} | "
            f"{record['raw_text']}"
        )

    if warnings:
        print()
        print("Warnings:")
        print("-" * 80)

        for warning in warnings:
            print(
                f"Row {warning['row']}, "
                f"column {warning['column']}: "
                f"{warning['value']} -> "
                f"{warning['message']}"
            )


if __name__ == "__main__":
    records, warnings = (
        extract_raw_timetable_records(
            r"data\Computing Undergraduate Timetable Fall Semester 2026.docx"
        )
    )

    print_summary(
        records=records,
        warnings=warnings,
    )